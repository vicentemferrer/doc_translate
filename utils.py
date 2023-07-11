import docx
import requests
import os

def translate(from_lang, to_lang, paragraphs):
    url = "https://lecto-translation.p.rapidapi.com/v1/translate/text"
    
    headers = {
        "content-type": "application/json",
        "X-RapidAPI-Key": "{key}",
        "X-RapidAPI-Host": "lecto-translation.p.rapidapi.com"
    }
    
    payload = {
		"texts": [*paragraphs],
        "to": [to_lang],
        "from": from_lang
	}

    response = requests.post(url, json=payload, headers=headers).json()
    return response['translations'][0]["translated"]

def get_paragraphs(path):
    doc = docx.Document(path)
    isolate_paragraphs = lambda paragraph: paragraph.text
    return list(map(isolate_paragraphs, doc.paragraphs))

def build_document(paragraphs, name):
    doc = docx.Document()
    doc_name = f"{name}.docx"
    add_paragraphs = lambda paragraph: doc.add_paragraph(paragraph)

    list(map(add_paragraphs, paragraphs))

    doc.save(doc_name)

    return f"{os.getcwd()}\\{doc_name}"